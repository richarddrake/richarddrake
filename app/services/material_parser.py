# 这个模块负责解析图片、文档、表格和文本材料，并提取可用于生成用例的上下文内容。
from __future__ import annotations

import csv
import io
import json
import re
from html.parser import HTMLParser
from pathlib import Path

from fastapi import HTTPException, UploadFile
from openpyxl import load_workbook

from app.schemas import UploadedMaterial


MAX_FILE_BYTES = 20 * 1024 * 1024
MAX_MATERIAL_COUNT = 12
MAX_TEXT_PER_FILE = 7000

TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".prd",
    ".json",
    ".yaml",
    ".yml",
    ".log",
    ".feature",
    ".html",
    ".htm",
}
SPREADSHEET_EXTENSIONS = {".xlsx", ".xlsm", ".xls", ".csv", ".tsv"}
WORD_EXTENSIONS = {".docx"}
PDF_EXTENSIONS = {".pdf"}


async def read_upload_materials(files: list[UploadFile]) -> list[UploadedMaterial]:
    materials: list[UploadedMaterial] = []
    if len(files) > MAX_MATERIAL_COUNT:
        raise HTTPException(status_code=400, detail=f"最多支持上传 {MAX_MATERIAL_COUNT} 个材料文件。")

    for file in files:
        if not file.filename:
            continue
        content_type = file.content_type or "application/octet-stream"
        data = await file.read()
        if len(data) > MAX_FILE_BYTES:
            raise HTTPException(status_code=413, detail=f"{file.filename} 超过 20MB。")
        materials.append(parse_material(file.filename, content_type, data))

    return materials


def parse_material(filename: str, content_type: str, data: bytes) -> UploadedMaterial:
    extension = Path(filename).suffix.lower()
    kind = _detect_kind(extension, content_type)
    extracted_text = ""
    note = ""

    try:
        if kind == "image":
            note = "图片将作为视觉材料发送给模型"
        elif kind == "spreadsheet":
            extracted_text = _extract_spreadsheet(extension, data)
        elif kind == "word":
            extracted_text = _extract_docx(data)
        elif kind == "pdf":
            extracted_text = _extract_pdf(data)
        elif kind == "text":
            extracted_text = _extract_text(extension, data)
        else:
            note = "暂不支持解析该文件内容，仅保留文件名作为来源线索"
    except Exception as exc:
        note = f"解析失败：{type(exc).__name__}"

    return UploadedMaterial(
        filename=filename,
        content_type=content_type,
        data=data,
        kind=kind,
        extracted_text=_limit_text(extracted_text),
        note=note,
    )


def build_material_context(materials: list[UploadedMaterial], references: str) -> str:
    lines: list[str] = []
    if references.strip():
        lines.append("外部文档/链接：")
        lines.append(_limit_text(references.strip(), 4000))

    text_materials = [item for item in materials if item.extracted_text]
    if text_materials:
        lines.append("可解析文件内容：")
        for item in text_materials:
            lines.append(f"\n【{item.filename}】")
            lines.append(item.extracted_text)

    notes = [item.describe() for item in materials if item.note and not item.extracted_text]
    if notes:
        lines.append("其他材料说明：")
        lines.extend(f"- {note}" for note in notes)

    return "\n".join(lines).strip()


def _detect_kind(extension: str, content_type: str) -> str:
    if content_type.startswith("image/"):
        return "image"
    if extension in SPREADSHEET_EXTENSIONS:
        return "spreadsheet"
    if extension in WORD_EXTENSIONS:
        return "word"
    if extension in PDF_EXTENSIONS:
        return "pdf"
    if extension in TEXT_EXTENSIONS or content_type.startswith("text/"):
        return "text"
    return "file"


def _extract_spreadsheet(extension: str, data: bytes) -> str:
    if extension in {".csv", ".tsv"}:
        return _extract_csv(data, delimiter="\t" if extension == ".tsv" else ",")
    if extension == ".xls":
        return _extract_xls(data)
    return _extract_xlsx(data)


def _extract_xlsx(data: bytes) -> str:
    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets[:8]:
        parts.append(f"# Sheet: {sheet.title}")
        row_count = 0
        for row in sheet.iter_rows(max_row=120, max_col=40, values_only=True):
            values = [_cell_to_text(value) for value in row]
            while values and not values[-1]:
                values.pop()
            if not any(values):
                continue
            parts.append(" | ".join(values))
            row_count += 1
            if row_count >= 80:
                parts.append("... 已截断更多行")
                break
    return "\n".join(parts)


def _extract_xls(data: bytes) -> str:
    try:
        import xlrd
    except ImportError:
        return "旧版 .xls 文件需要安装 xlrd 后才能解析。建议另存为 .xlsx 后上传。"

    workbook = xlrd.open_workbook(file_contents=data)
    parts: list[str] = []
    for sheet in workbook.sheets()[:8]:
        parts.append(f"# Sheet: {sheet.name}")
        for row_index in range(min(sheet.nrows, 80)):
            values = [_cell_to_text(sheet.cell_value(row_index, col_index)) for col_index in range(min(sheet.ncols, 40))]
            while values and not values[-1]:
                values.pop()
            if values:
                parts.append(" | ".join(values))
        if sheet.nrows > 80:
            parts.append("... 已截断更多行")
    return "\n".join(parts)


def _extract_csv(data: bytes, delimiter: str) -> str:
    text = _decode_bytes(data)
    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    rows: list[str] = []
    for index, row in enumerate(reader):
        if index >= 120:
            rows.append("... 已截断更多行")
            break
        values = [item.strip() for item in row]
        if any(values):
            rows.append(" | ".join(values[:40]))
    return "\n".join(rows)


def _extract_text(extension: str, data: bytes) -> str:
    text = _decode_bytes(data)
    if extension in {".html", ".htm"}:
        parser = _TextHTMLParser()
        parser.feed(text)
        return parser.text()
    if extension == ".json":
        try:
            parsed = json.loads(text)
            return json.dumps(parsed, ensure_ascii=False, indent=2)
        except json.JSONDecodeError:
            return text
    return text


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        return "Word .docx 文件需要安装 python-docx 后才能解析。"

    document = Document(io.BytesIO(data))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables[:10]:
        for row in table.rows[:80]:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(values):
                parts.append(" | ".join(values))
    return "\n".join(parts)


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "PDF 文件需要安装 pypdf 后才能解析。"

    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for index, page in enumerate(reader.pages[:12], 1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"# Page {index}\n{text.strip()}")
    if len(reader.pages) > 12:
        pages.append("... 已截断更多页")
    return "\n\n".join(pages)


def _decode_bytes(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _cell_to_text(value: object) -> str:
    if value is None:
        return ""
    text = str(value).strip()
    return re.sub(r"\s+", " ", text)


def _limit_text(text: str, limit: int = MAX_TEXT_PER_FILE) -> str:
    text = "\n".join(line.rstrip() for line in text.splitlines())
    text = re.sub(r"\n{4,}", "\n\n\n", text).strip()
    if len(text) <= limit:
        return text
    return text[:limit].rstrip() + "\n... 内容已截断"


class _TextHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []

    def handle_data(self, data: str) -> None:
        text = data.strip()
        if text:
            self._parts.append(text)

    def text(self) -> str:
        return "\n".join(self._parts)
